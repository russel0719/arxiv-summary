from langchain_perplexity import ChatPerplexity
from langchain_core.messages import HumanMessage, SystemMessage

import os
from dotenv import load_dotenv
import re
import arxiv
from pydantic import BaseModel, Field
from typing import List
from datetime import datetime
from docx import Document

load_dotenv()
PPLX_API_KEY = os.getenv("PPLX_API_KEY")
EXCEPT_KEYWORDS = ["robot", "security", "cripto", "emotion", "study", "report", "survey", "review"]
BATCH_SIZE = 25
MAX_SIZE = 1000

client = arxiv.Client(page_size=100, delay_seconds=3, num_retries=5)
DOCUMENT_PATH = os.path.join(os.path.dirname(__file__), "{year}/{month}/{day}_{section}.docx")


class Answerformat(BaseModel):
    korean_summaries: List[str] = Field(description="List of Korean summaries for each paper abstract.")

    
def get_papers(section: str, last_submitted_date: str=None, max_results: int=BATCH_SIZE) -> List[arxiv.Result]:
    last_submitted_date = str(int(last_submitted_date) + 1) if last_submitted_date else None
    now = datetime.now().strftime("%Y%m%d%H%M")
    except_query = " OR ".join([f"ti:{keyword}" for keyword in EXCEPT_KEYWORDS])
    default_query = f"cat:{section} AND submittedDate:[{last_submitted_date} TO {now}]" if last_submitted_date else f"cat:{section}"
    search = arxiv.Search(
        query=f"{default_query} ANDNOT ({except_query})",
        max_results=max_results,
        sort_by=arxiv.SortCriterion.SubmittedDate,
        sort_order=arxiv.SortOrder.Ascending,
    )
    results = list(client.results(search))
    return results


def get_summary_from_abstract(papers: List[arxiv.Result]) -> List[dict]:
    summaries = []
    for paper in papers:
        summary = {"title": paper.title, "url":paper.entry_id, "summary": paper.summary, "submitted_date": paper.published}
        summaries.append(summary)
    return summaries


def add_translations(llm: ChatPerplexity, summaries: List[dict]) -> List[dict]:
    formatted_abstracts = []
    for idx, s in enumerate(summaries, start=1):
        formatted_abstracts.append(f"{idx}. {s['summary'].strip().replace('\n', ' ')}")

    num_summaries = len(formatted_abstracts)
    prompt = (
        f"You will receive a numbered list of {num_summaries} paper abstracts. Provide korean summaries for each of the abstracts of ALL papers. Each summary should be important thoughts of the abstract.\n"
        f"Return your answer as JSON format with the key 'korean_summaries' and the value as a list of {num_summaries} korean summaries.\n"
        "For example, if you receive 3 abstracts, your answer should be like this:\n\n"
        '{"korean_summaries": ["번역내용1", "번역내용2", "번역내용3"]}\n\n'
        "Do not skip any paper's abstract translation in answer. Maintain the input order strictly. Do not use any LaTeX symbol and XML tags in your answer.\n"
    )

    messages = [
        SystemMessage(content=prompt),
        HumanMessage(content="\n\n".join(formatted_abstracts))
    ]

    ai_message = llm.invoke(messages) # type: Answerformat
    translations = ai_message.korean_summaries
    translations = [re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', translation.replace("$", "")).strip() for translation in translations]

    if len(translations) != len(summaries):
        print(f"[Warning] Mismatch in count: {len(translations)} translations for {len(summaries)} summaries.")
        print(f"[Warning] {len(translations) - 1} translations will be applied.")

    for i, summary in enumerate(summaries):
        print(f"[{i + 1}] {summary['title']}:")
        print(translations[i] if i < len(translations) else "[번역 누락]")
        print()
        summary["korean_summary"] = translations[i] if i < len(translations) else "[번역 누락]"
    
    last_idx = len(translations) - 1 if len(translations) < len(summaries) and len(translations) > 0 else len(translations)
    return summaries[:last_idx]


def write_document(section: str, summary: dict) -> None:
    print(f"Writing summary for `{summary['title']}` in {section} section")
    date = summary["submitted_date"].strftime("%Y%m%d")
    year, month, day = date[:4], date[4:6], date[6:8]
    document_path = DOCUMENT_PATH.format(year=year, month=month, day=day, section=section)
    os.makedirs(os.path.dirname(document_path), exist_ok=True)
    if not os.path.exists(document_path):
        doc = Document()
        doc.add_heading(f"{date} Papers Summary in {section} section", level=0)
        doc.save(document_path)
    doc = Document(document_path)
    doc.add_heading(summary["title"], level=2)
    p = doc.add_paragraph()
    p.add_run("URL: ").bold = True
    p.add_run(summary["url"])
    p.add_run("\n")
    p.add_run("Summary: ").bold = True
    p.add_run(summary["summary"])
    p.add_run("\n")
    p.add_run("Korean Summary: ").bold = True
    p.add_run(summary["korean_summary"])
    p.add_run("\n")
    doc.save(document_path)


def write_document_with_latest_papers(section: str, llm: ChatPerplexity):
    if os.path.exists("last_submitted_date.txt"):
        with open("last_submitted_date.txt", "r") as f:
            last_submitted_date = f.read().strip()
    else:
        last_submitted_date = None
    
    papers = get_papers(section, last_submitted_date)
    print(f"Fetched {len(papers)} papers from {section} section.")
    if len(papers) < BATCH_SIZE:
        return 0
    summaries = get_summary_from_abstract(papers)

    if summaries:
        summaries = add_translations(llm, summaries)
        for summary in summaries:
            write_document(section, summary)
    
    # Update the last submitted date
    if summaries:
        with open("last_submitted_date.txt", "w") as f:
            last_submitted_date = summaries[-1]["submitted_date"].strftime("%Y%m%d%H%M")
            f.write(last_submitted_date)
    return len(summaries)

    
if __name__ == "__main__":
    # Initialize the LLM and agent
    llm = ChatPerplexity(model="sonar", api_key=PPLX_API_KEY, temperature=0.0)
    structed_llm = llm.with_structured_output(Answerformat)

    num_papers = input(f"How many papers do you want to fetch? (default: max): ")
    if not num_papers:
        num_papers = MAX_SIZE
    else:
        num_papers = int(num_papers)
    total_cnt = 0
    while total_cnt < num_papers:
        cnt = write_document_with_latest_papers("cs.CV", structed_llm)
        total_cnt += cnt
        if cnt == 0:
            print("No more papers to fetch.")
            break
    print(f"Total {total_cnt} papers are summarized.")